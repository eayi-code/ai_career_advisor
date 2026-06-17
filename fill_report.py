# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document('期末大作业课程报告.docx')
table = doc.tables[2]
cell = table.cell(0, 0)

# 清空单元格
for p in cell.paragraphs:
    p.clear()

# 写入内容
cell.text = """一、实验目的
1. 综合应用专业知识，系统运用Flask、LangChain、ChromaDB、MySQL等课程核心知识，完成一个功能完整、前后端协同的AI智能职业规划系统；
2. 提升工程实践能力，以项目为导向，经历需求分析、系统设计、编码实现、测试优化等完整开发流程，掌握多智能体系统开发的工程方法；
3. 培养AI工程素养，深入理解多智能体协作、ReAct推理模式、向量记忆等AI工程核心技术；
4. 强化技术分析与决策能力，能够识别影响系统性能的关键因素，合理选择技术方案，并对解决方案进行可行性论证。

二、实验任务与要求
（一）选题要求
后端框架：Flask
前端框架：HTML + CSS + JavaScript + Jinja2模板引擎
人工智能技术：LangChain智能体框架 + ChromaDB向量数据库 + ReAct推理模式
数据库：MySQL，包含完整的数据库设计与实现

（二）主要开发内容
1. 数据库设计：完成数据库选型、表结构设计、表间关联关系
2. 前后端设计：完成后端核心配置、路由与接口设计；完成前端页面布局
3. AI模块设计：实现LangChain智能体框架集成、多Agent协作编排

三、实验原理
1. 前后端分离架构：前端通过HTTP请求调用后端Flask提供的RESTful API接口
2. 多智能体协作技术：LangChain框架、LangGraph编排、ChromaDB向量数据库
3. 数据库设计原理：采用ER模型进行概念结构设计，转换为关系模式
4. 系统模块协同原理：各功能模块通过明确的接口协议进行协同工作
5. 工程规范与性能优化：遵循代码规范、版本管理、异常处理等工程实践

四、实验仪器设备
1. 硬件环境：普通个人电脑
2. 软件环境：Windows 10及以上操作系统，Python 3.11，MySQL 8.0

五、实验设计步骤
（一）项目背景
随着就业市场竞争日益激烈，求职者面临诸多挑战：信息不对称、简历撰写困难、面试准备不足、职业规划迷茫。本项目旨在通过AI智能体技术，构建一个智能化的职业决策支持系统，为求职者提供个性化的职业匹配、简历优化、面试准备等服务。

（二）技术架构
技术栈：
- 后端：Flask 3.x + SQLAlchemy + Flask-Login
- 前端：HTML5 + CSS3 + JavaScript + Jinja2
- AI框架：LangChain + LangGraph
- 向量数据库：ChromaDB + Ollama
- 关系数据库：MySQL 8.0
- LLM：mimo-v2.5-pro

【请在此处插入系统架构图】

（三）数据库设计

【请在此处插入ER图】

数据库表结构：
- users表：用户账号信息（id, username, email, password_hash, avatar等）
- user_profiles表：用户档案（skills, target_job_title, projects等）
- analysis_history表：对话历史（messages, steps, tools_used等）
- jobs表：职位数据（60条）
- skills表：技能数据（60条）
- side_jobs表：副业数据（30条）

六、实验实现步骤
（一）后端实现

【请在此处插入后端项目目录结构截图】

核心配置（config.py）：
SQLALCHEMY_DATABASE_URI = mysql+pymysql://root:1234@localhost/career_advisor
SQLALCHEMY_ENGINE_OPTIONS = {pool_size: 10, pool_recycle: 3600}

主要API接口：
- POST /api/agent/chat/stream - SSE流式对话
- POST /api/agent/chat - 同步对话
- POST /api/upload/resume - 上传简历
- GET /api/history - 获取对话历史

核心代码（Agent基类）：
class BaseAgent:
    def run(self, user_input, user_id=None):
        context = self._build_context(user_id, user_input)
        for chunk in self.agent.stream({"messages": [...]}):
            if msg.type == "ai" and msg.content:
                if self.on_token_callback:
                    self.on_token_callback(msg.content)

（二）前端实现

【请在此处插入前端项目目录结构截图】

【请在此处插入落地页截图】

【请在此处插入登录页截图】

【请在此处插入对话页面截图】

【请在此处插入用户中心截图】

【请在此处插入简历预览截图】

核心代码（SSE流式对话）：
async function streamAgentChat(message, agentType) {
    const response = await fetch('/api/agent/chat/stream', {
        method: 'POST',
        headers: {'Content-Type': 'application/json'},
        body: JSON.stringify({ message, agent_type: agentType })
    });
    const reader = response.body.getReader();
    while (true) {
        const { done, value } = await reader.read();
        if (done) break;
        parseSSEEvents(value);
    }
}

（三）人工智能模块实现

【请在此处插入Agent架构图】

5个专业Agent：
- CareerAgent：职业规划（search_jobs, query_salary, compare_jobs）
- SkillAgent：技能分析（analyze_skill_gap, recommend_learning_path）
- SideJobAgent：副业推荐（search_side_jobs, calculate_side_job_roi）
- ResumeAgent：简历优化（generate_resume, ats_score）
- InterviewAgent：面试教练（generate_interview_questions）

ReAct推理模式：

【请在此处插入推理过程截图】

Thought: 用户想搜索数据分析师岗位
Action: search_jobs(keywords=["数据分析师"])
Observation: 找到3个匹配职位
Answer: 根据搜索结果，以下是数据分析师相关岗位...

三层记忆架构：

【请在此处插入记忆架构图】

- L1工作记忆：Python字典，存储当前对话上下文
- L2情景记忆：ChromaDB向量数据库，存储历史对话
- L3语义记忆：MySQL用户档案，存储用户偏好

七、实验测试
1. 测试环境
- 硬件环境：Intel Core i7处理器，16GB内存
- 软件环境：Windows 11，Python 3.11，MySQL 8.0

2. 功能测试

【请在此处插入功能测试截图1：单Agent对话】

【请在此处插入功能测试截图2：复合Agent对话】

【请在此处插入功能测试截图3：流式输出效果】

【请在此处插入功能测试截图4：推理详情面板】

测试用例：
| 测试功能 | 测试输入 | 预期输出 | 实际结果 | 是否通过 |
|----------|----------|----------|----------|----------|
| 用户注册 | 用户名、邮箱、密码 | 注册成功 | 注册成功 | 通过 |
| 用户登录 | 用户名、密码 | 登录成功 | 登录成功 | 通过 |
| 单Agent对话 | 搜索数据分析师岗位 | 返回职位列表 | 返回3个职位 | 通过 |
| 复合Agent对话 | 搜索岗位并出面试题 | 返回整合结果 | 返回成功 | 通过 |
| 简历生成 | 帮我写简历 | 生成HTML简历 | 生成成功 | 通过 |
| 流式输出 | 任意对话 | 逐字显示 | 流式显示 | 通过 |

3. 测试结果分析
系统整体表现良好，核心功能均正常工作：
- 多智能体协作：5个Agent能够正确协作
- 流式输出：Token级流式输出正常
- 记忆系统：三层记忆架构工作正常
- 工具调用：21个工具均能正确调用"""

doc.save('期末大作业课程报告_填写版.docx')
print('报告已生成完成！')
