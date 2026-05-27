from flask import Blueprint, request, jsonify, send_file, Response, stream_with_context, current_app
from flask_login import login_required, current_user
from datetime import datetime
import uuid
import os
import io
import base64
import json
import queue
import threading
from werkzeug.utils import secure_filename
from sqlalchemy.orm.attributes import flag_modified
from app import db
from app.models.user import User
from app.models.history import AnalysisHistory
from app.agents.orchestrator import AgentOrchestrator

api_bp = Blueprint('api', __name__)
orchestrator = AgentOrchestrator()

# 任务存储（使用模块级变量）
task_store = {}

# SSE任务队列存储
sse_task_queues = {}


@api_bp.route('/test', methods=['GET'])
def test_route():
    return jsonify({"code": 200, "message": "API is working"})


@api_bp.route('/agent/chat/async', methods=['POST'])
@login_required
def agent_chat_async():
    """异步Agent对话（轮询模式，支持实时推理步骤）"""
    data = request.get_json()
    message = data.get('message', '')
    agent_type = data.get('agent_type', 'auto')
    conversation_id = data.get('conversation_id')

    if not message:
        return jsonify({"code": 400, "error": "消息不能为空"}), 400

    if not conversation_id:
        conversation_id = str(uuid.uuid4())[:8]

    force_agent = agent_type if agent_type != 'auto' else None

    # 生成任务ID
    task_id = str(uuid.uuid4())[:12]
    user_id = current_user.id
    
    # 获取Flask应用对象（用于在线程中访问数据库）
    from flask import current_app
    app = current_app._get_current_object()
    
    # 存储任务状态
    task_store[task_id] = {
        'status': 'pending',
        'progress': '正在准备...',
        'steps': [],  # 实时推理步骤
        'result': None,
        'error': None,
        'conversation_id': conversation_id
    }
    
    # 后台执行Agent
    import threading
    def run_agent():
        with app.app_context():
            try:
                task_store[task_id]['status'] = 'running'
                task_store[task_id]['progress'] = '正在分析您的问题...'
                
                # 使用自定义orchestrator，支持实时步骤更新
                result = _run_orchestrator_with_steps(message, user_id, force_agent, task_id)
                
                task_store[task_id]['status'] = 'completed'
                task_store[task_id]['result'] = result
                task_store[task_id]['progress'] = '处理完成'
                    
            except Exception as e:
                import traceback
                traceback.print_exc()
                task_store[task_id]['status'] = 'failed'
                task_store[task_id]['error'] = str(e)
    
    thread = threading.Thread(target=run_agent)
    thread.start()
    
    return jsonify({
        "code": 200,
        "data": {
            "task_id": task_id,
            "conversation_id": conversation_id,
            "status": "pending"
        }
    })


def _run_orchestrator_with_steps(message, user_id, force_agent, task_id):
    """执行orchestrator并实时更新步骤（模拟实时进度）"""
    import time
    from app.agents.orchestrator import AgentOrchestrator
    
    # 模拟步骤进度
    simulated_steps = [
        {
            "type": "intent_analysis",
            "title": "正在分析您的意图",
            "detail": "识别问题类型和关键词",
            "status": "running"
        }
    ]
    task_store[task_id]['steps'] = simulated_steps
    task_store[task_id]['progress'] = '正在分析意图...'
    
    time.sleep(0.5)  # 模拟分析时间
    
    # 确定使用哪个Agent
    agent_name = force_agent or "career"
    agent_names = {
        "career": "职业规划顾问",
        "skill": "技能发展顾问",
        "side_job": "副业规划专家",
        "resume": "简历优化专家",
        "interview": "面试教练"
    }
    
    # 更新为分析完成，准备调用Agent
    simulated_steps[0]["status"] = "completed"
    simulated_steps[0]["detail"] = f"识别为{agent_names.get(agent_name, agent_name)}任务"
    simulated_steps.append({
        "type": "agent_call",
        "title": f"调用{agent_names.get(agent_name, agent_name)}",
        "detail": "正在准备执行任务",
        "status": "running"
    })
    task_store[task_id]['steps'] = simulated_steps
    task_store[task_id]['progress'] = f'正在调用{agent_names.get(agent_name, agent_name)}...'
    
    time.sleep(0.3)
    
    # 更新Agent调用状态
    simulated_steps[1]["detail"] = "正在搜索职位数据..."
    simulated_steps.append({
        "type": "tool",
        "title": "执行工具调用",
        "detail": "search_jobs - 搜索相关职位",
        "status": "running"
    })
    task_store[task_id]['steps'] = simulated_steps
    task_store[task_id]['progress'] = '正在搜索数据...'
    
    # 执行实际的orchestrator
    orch = AgentOrchestrator()
    result = orch.process(message, user_id, force_agent)
    
    # 执行完成后，用真实步骤替换模拟步骤
    if result.get("success"):
        execution_steps = result.get("steps", [])
        intermediate_steps = result.get("intermediate_steps", [])
        
        # 合并所有步骤
        all_steps = []
        if execution_steps:
            all_steps.extend(execution_steps)
        if intermediate_steps:
            all_steps.extend(intermediate_steps)
        
        # 如果没有步骤，使用模拟步骤但标记为完成
        if not all_steps:
            for step in simulated_steps:
                step["status"] = "completed"
            all_steps = simulated_steps
        
        task_store[task_id]['steps'] = all_steps
        task_store[task_id]['progress'] = f'已完成 {len(all_steps)} 个步骤'
    
    return result


@api_bp.route('/agent/task/<task_id>', methods=['GET'])
@login_required
def get_task_status(task_id):
    """查询任务状态（包含实时推理步骤）"""
    if task_id not in task_store:
        return jsonify({"code": 404, "error": "任务不存在"}), 404
    
    task = task_store[task_id]
    
    response = {
        "task_id": task_id,
        "status": task['status'],
        "progress": task['progress'],
        "steps": task.get('steps', [])  # 返回实时步骤
    }
    
    if task['status'] == 'completed' and task['result']:
        response['result'] = task['result']
        # 任务完成后清理
        # del task_store[task_id]  # 暂时不清理，方便调试
    elif task['status'] == 'failed' and task['error']:
        response['error'] = task['error']
    
    return jsonify({"code": 200, "data": response})


@api_bp.route('/agent/chat', methods=['POST'])
@login_required
def agent_chat():
    data = request.get_json()
    message = data.get('message', '')
    agent_type = data.get('agent_type', 'auto')
    conversation_id = data.get('conversation_id')

    if not message:
        return jsonify({"code": 400, "error": "消息不能为空"}), 400

    if not conversation_id:
        conversation_id = str(uuid.uuid4())[:8]

    force_agent = agent_type if agent_type != 'auto' else None
    
    try:
        result = orchestrator.process(message, current_user.id, force_agent)
    except Exception as e:
        db.session.rollback()
        return jsonify({"code": 500, "error": f"Agent执行失败: {str(e)}"}), 500

    if result["success"]:
        try:
            history = AnalysisHistory.query.filter_by(
                conversation_id=conversation_id,
                user_id=current_user.id
            ).first()

            if not history:
                title = message[:30] + ('...' if len(message) > 30 else '')
                history = AnalysisHistory(
                    user_id=current_user.id,
                    conversation_id=conversation_id,
                    title=title,
                    analysis_type=agent_type,
                    agent_used=result.get("agent_used", "unknown"),
                    input_data={"message": message},
                    result_data={"output": result.get("output", "")},
                    reasoning_steps=result.get("intermediate_steps", []),
                    tools_used=[s["action"] for s in result.get("intermediate_steps", [])],
                    messages=[]
                )
                db.session.add(history)
            else:
                history.input_data = {"message": message}
                history.result_data = {"output": result.get("output", "")}
                history.reasoning_steps = result.get("intermediate_steps", [])
                existing_tools = history.tools_used or []
                new_tools = [s["action"] for s in result.get("intermediate_steps", [])]
                history.tools_used = list(set(existing_tools + new_tools))
                history.agent_used = result.get("agent_used", history.agent_used)

            messages = history.messages or []
            messages.append({
                "role": "user",
                "content": message,
                "timestamp": datetime.utcnow().isoformat()
            })
            messages.append({
                "role": "assistant",
                "content": result.get("output", ""),
                "agent": result.get("agent_used"),
                "steps": result.get("intermediate_steps", []),
                "timestamp": datetime.utcnow().isoformat()
            })
            history.messages = messages
            history.updated_at = datetime.utcnow()
            
            # 标记JSON字段为已修改，确保SQLAlchemy检测到变化
            flag_modified(history, 'messages')
            flag_modified(history, 'tools_used')

            db.session.commit()
        except Exception as e:
            db.session.rollback()
            print(f"保存历史记录失败: {e}")
            # 不影响返回结果，继续执行

        return jsonify({
            "code": 200,
            "data": {
                "conversation_id": conversation_id,
                "response": result.get("output", ""),
                "agent_used": result.get("agent_used"),
                "reasoning_steps": result.get("intermediate_steps", []),
                "tools_used": list(set(s["action"] for s in result.get("intermediate_steps", []))),
                "execution_steps": result.get("steps", []),
                "is_composite": result.get("is_composite", False)
            }
        })
    else:
        return jsonify({"code": 500, "error": result.get("error", "处理失败")}), 500


@api_bp.route('/history/<conversation_id>', methods=['GET'])
@login_required
def get_conversation(conversation_id):
    # 获取该conversation_id的最新记录
    history = AnalysisHistory.query.filter_by(
        conversation_id=conversation_id,
        user_id=current_user.id
    ).order_by(AnalysisHistory.updated_at.desc()).first()

    if not history:
        return jsonify({"code": 404, "error": "对话不存在"}), 404

    return jsonify({
        "code": 200,
        "data": {
            "conversation_id": history.conversation_id,
            "title": history.title,
            "agent_used": history.agent_used,
            "messages": history.messages or [],
            "created_at": history.created_at.isoformat(),
            "updated_at": history.updated_at.isoformat() if history.updated_at else None
        }
    })


@api_bp.route('/history', methods=['GET'])
@login_required
def get_history():
    # 获取每个conversation_id的最新记录
    from sqlalchemy import func
    
    # 先获取每个conversation_id的最大updated_at
    subquery = db.session.query(
        AnalysisHistory.conversation_id,
        func.max(AnalysisHistory.updated_at).label('max_updated')
    ).filter_by(user_id=current_user.id).group_by(
        AnalysisHistory.conversation_id
    ).subquery()
    
    # 获取完整的记录
    histories = db.session.query(AnalysisHistory).join(
        subquery,
        db.and_(
            AnalysisHistory.conversation_id == subquery.c.conversation_id,
            AnalysisHistory.updated_at == subquery.c.max_updated
        )
    ).filter(
        AnalysisHistory.user_id == current_user.id
    ).order_by(
        AnalysisHistory.updated_at.desc()
    ).limit(50).all()

    return jsonify({
        "code": 200,
        "data": [{
            "conversation_id": h.conversation_id,
            "title": h.title,
            "agent_used": h.agent_used,
            "message_count": len(h.messages) if h.messages else 0,
            "last_message": h.messages[-1]["content"][:50] if h.messages else "",
            "created_at": h.created_at.isoformat(),
            "updated_at": h.updated_at.isoformat() if h.updated_at else None
        } for h in histories]
    })


@api_bp.route('/history/<conversation_id>', methods=['DELETE'])
@login_required
def delete_conversation(conversation_id):
    try:
        # 删除该conversation_id的所有记录
        count = AnalysisHistory.query.filter_by(
            conversation_id=conversation_id,
            user_id=current_user.id
        ).delete()
        db.session.commit()
        return jsonify({"code": 200, "message": "已删除", "count": count})
    except Exception as e:
        db.session.rollback()
        return jsonify({"code": 500, "error": str(e)}), 500


@api_bp.route('/history/save', methods=['POST'])
@login_required
def save_conversation():
    """保存对话记录（用于异步模式）"""
    data = request.get_json()
    conversation_id = data.get('conversation_id')
    message = data.get('message', '')
    result = data.get('result', {})
    
    if not conversation_id or not message:
        return jsonify({"code": 400, "error": "参数不完整"}), 400
    
    try:
        history = AnalysisHistory.query.filter_by(
            conversation_id=conversation_id,
            user_id=current_user.id
        ).first()
        
        if not history:
            title = message[:30] + ('...' if len(message) > 30 else '')
            history = AnalysisHistory(
                user_id=current_user.id,
                conversation_id=conversation_id,
                title=title,
                analysis_type='auto',
                agent_used=result.get("agent_used", "unknown"),
                input_data={"message": message},
                result_data={"output": result.get("output", "")},
                reasoning_steps=result.get("intermediate_steps", []),
                tools_used=[s.get("action", "") for s in result.get("intermediate_steps", [])],
                messages=[]
            )
            db.session.add(history)
        else:
            history.input_data = {"message": message}
            history.result_data = {"output": result.get("output", "")}
            history.reasoning_steps = result.get("intermediate_steps", [])
            existing_tools = history.tools_used or []
            new_tools = [s.get("action", "") for s in result.get("intermediate_steps", [])]
            history.tools_used = list(set(existing_tools + new_tools))
            history.agent_used = result.get("agent_used", history.agent_used)
        
        messages = history.messages or []
        messages.append({
            "role": "user",
            "content": message,
            "timestamp": datetime.utcnow().isoformat()
        })
        messages.append({
            "role": "assistant",
            "content": result.get("output", ""),
            "agent": result.get("agent_used"),
            "steps": result.get("intermediate_steps", []),
            "execution_steps": result.get("steps", []),
            "timestamp": datetime.utcnow().isoformat()
        })
        history.messages = messages
        history.updated_at = datetime.utcnow()
        
        # 标记JSON字段为已修改
        flag_modified(history, 'messages')
        flag_modified(history, 'tools_used')
        
        db.session.commit()
        return jsonify({"code": 200, "message": "保存成功"})
        
    except Exception as e:
        db.session.rollback()
        import traceback
        traceback.print_exc()
        return jsonify({"code": 500, "error": str(e)}), 500


@api_bp.route('/agent/chat/stream', methods=['POST'])
@login_required
def agent_chat_stream():
    """SSE流式响应 - 实时推理步骤"""
    data = request.get_json()
    message = data.get('message', '')
    agent_type = data.get('agent_type', 'auto')
    conversation_id = data.get('conversation_id')
    last_agent = data.get('last_agent')  # 上一次使用的Agent

    if not message:
        return jsonify({"code": 400, "error": "消息不能为空"}), 400

    if not conversation_id:
        conversation_id = str(uuid.uuid4())[:8]

    force_agent = agent_type if agent_type != 'auto' else None
    user_id = current_user.id
    task_id = str(uuid.uuid4())[:12]
    
    # 创建任务队列
    task_queue = queue.Queue()
    sse_task_queues[task_id] = task_queue
    
    # 获取Flask应用对象
    app = current_app._get_current_object()
    
    # 获取对话历史（用于上下文理解）
    conversation_history = []
    try:
        db.session.rollback()
        histories = AnalysisHistory.query.filter_by(
            conversation_id=conversation_id,
            user_id=user_id
        ).order_by(AnalysisHistory.updated_at.desc()).first()
        if histories and histories.messages:
            conversation_history = histories.messages[-10:]  # 最近10条消息
    except Exception as e:
        print(f"获取对话历史失败: {e}")
    
    def on_step_callback(step_data):
        """步骤更新回调函数（在后台线程中被调用）"""
        task_queue.put({
            'type': 'step',
            'data': step_data
        })
    
    def generate():
        """SSE生成器"""
        try:
            # 发送开始事件
            yield f"event: start\ndata: {json.dumps({'task_id': task_id, 'conversation_id': conversation_id})}\n\n"
            
            # 等待任务完成
            while True:
                try:
                    event = task_queue.get(timeout=120)  # 120秒超时
                    event_type = event.get('type')
                    
                    if event_type == 'step':
                        # 发送步骤更新
                        yield f"event: step\ndata: {json.dumps(event['data'])}\n\n"
                    
                    elif event_type == 'progress':
                        # 发送进度更新
                        yield f"event: progress\ndata: {json.dumps({'message': event['message']})}\n\n"
                    
                    elif event_type == 'done':
                        # 发送完成事件
                        yield f"event: done\ndata: {json.dumps(event['result'])}\n\n"
                        break
                    
                    elif event_type == 'error':
                        # 发送错误事件
                        yield f"event: error\ndata: {json.dumps({'error': event['error']})}\n\n"
                        break
                        
                except queue.Empty:
                    # 超时
                    yield f"event: error\ndata: {json.dumps({'error': '请求超时，请稍后重试'})}\n\n"
                    break
                    
        finally:
            # 清理队列
            if task_id in sse_task_queues:
                del sse_task_queues[task_id]
    
    def run_agent():
        """后台执行Agent"""
        with app.app_context():
            try:
                # 发送进度
                task_queue.put({'type': 'progress', 'message': '正在分析意图...'})
                
                # 创建带回调的orchestrator
                from app.agents.orchestrator import AgentOrchestrator
                orch = AgentOrchestrator(on_step_callback=on_step_callback)
                
                # 执行orchestrator（传递last_agent和conversation_history）
                result = orch.process(
                    message, 
                    user_id, 
                    force_agent,
                    last_agent=last_agent,
                    conversation_history=conversation_history
                )
                
                # 保存对话记录
                if result.get("success"):
                    try:
                        history = AnalysisHistory.query.filter_by(
                            conversation_id=conversation_id,
                            user_id=user_id
                        ).first()
                        
                        if not history:
                            title = message[:30] + ('...' if len(message) > 30 else '')
                            history = AnalysisHistory(
                                user_id=user_id,
                                conversation_id=conversation_id,
                                title=title,
                                analysis_type=agent_type,
                                agent_used=result.get("agent_used", "unknown"),
                                input_data={"message": message},
                                result_data={"output": result.get("output", "")},
                                reasoning_steps=result.get("intermediate_steps", []),
                                tools_used=[s.get("action", "") for s in result.get("intermediate_steps", [])],
                                messages=[]
                            )
                            db.session.add(history)
                        
                        messages = history.messages or []
                        messages.append({
                            "role": "user",
                            "content": message,
                            "timestamp": datetime.utcnow().isoformat()
                        })
                        messages.append({
                            "role": "assistant",
                            "content": result.get("output", ""),
                            "agent": result.get("agent_used"),
                            "steps": result.get("intermediate_steps", []),
                            "execution_steps": result.get("steps", []),
                            "timestamp": datetime.utcnow().isoformat()
                        })
                        history.messages = messages
                        history.updated_at = datetime.utcnow()
                        
                        flag_modified(history, 'messages')
                        flag_modified(history, 'tools_used')
                        
                        db.session.commit()
                        print(f"[SSE] 对话已保存: {conversation_id}")
                    except Exception as e:
                        db.session.rollback()
                        print(f"[SSE] 保存对话失败: {e}")
                
                # 发送完成事件
                if result.get("success"):
                    task_queue.put({
                        'type': 'done',
                        'result': result
                    })
                else:
                    task_queue.put({
                        'type': 'error',
                        'error': result.get('error', '执行失败')
                    })
                    
            except Exception as e:
                import traceback
                traceback.print_exc()
                task_queue.put({
                    'type': 'error',
                    'error': str(e)
                })
    
    # 启动后台线程
    thread = threading.Thread(target=run_agent)
    thread.start()
    
    # 返回SSE响应
    return Response(
        stream_with_context(generate()),
        mimetype='text/event-stream',
        headers={
            'Cache-Control': 'no-cache',
            'X-Accel-Buffering': 'no',
            'Connection': 'keep-alive'
        }
    )


@api_bp.route('/agent/tools', methods=['GET'])
@login_required
def agent_tools():
    status = orchestrator.get_agent_status()
    return jsonify({"code": 200, "data": status})


@api_bp.route('/user/update-username', methods=['POST'])
@login_required
def update_username():
    data = request.get_json()
    new_username = data.get('username', '').strip()
    
    if not new_username:
        return jsonify({"code": 400, "error": "用户名不能为空"}), 400
    
    if len(new_username) < 2 or len(new_username) > 50:
        return jsonify({"code": 400, "error": "用户名长度需在2-50个字符之间"}), 400
    
    existing = User.query.filter(User.username == new_username, User.id != current_user.id).first()
    if existing:
        return jsonify({"code": 400, "error": "用户名已被占用"}), 400
    
    try:
        current_user.username = new_username
        db.session.commit()
        return jsonify({"code": 200, "message": "用户名修改成功"})
    except Exception as e:
        db.session.rollback()
        return jsonify({"code": 500, "error": str(e)}), 500


@api_bp.route('/user/update-password', methods=['POST'])
@login_required
def update_password():
    data = request.get_json()
    old_password = data.get('old_password', '')
    new_password = data.get('new_password', '')
    
    if not old_password or not new_password:
        return jsonify({"code": 400, "error": "请填写完整"}), 400
    
    if not current_user.check_password(old_password):
        return jsonify({"code": 400, "error": "当前密码错误"}), 400
    
    if len(new_password) < 6:
        return jsonify({"code": 400, "error": "新密码长度至少6位"}), 400
    
    try:
        current_user.set_password(new_password)
        db.session.commit()
        return jsonify({"code": 200, "message": "密码修改成功"})
    except Exception as e:
        db.session.rollback()
        return jsonify({"code": 500, "error": str(e)}), 500


@api_bp.route('/user/update-avatar', methods=['POST'])
@login_required
def update_avatar():
    data = request.get_json()
    avatar_data = data.get('avatar', '')
    
    if not avatar_data:
        return jsonify({"code": 400, "error": "请选择头像"}), 400
    
    try:
        if avatar_data.startswith('data:image'):
            header, encoded = avatar_data.split(',', 1)
            image_data = base64.b64decode(encoded)
            
            upload_dir = os.path.join('app', 'static', 'uploads', 'avatars')
            os.makedirs(upload_dir, exist_ok=True)
            
            filename = f"avatar_{current_user.id}.png"
            filepath = os.path.join(upload_dir, filename)
            
            with open(filepath, 'wb') as f:
                f.write(image_data)
            
            current_user.avatar = f"/static/uploads/avatars/{filename}"
            db.session.commit()
            
            return jsonify({"code": 200, "message": "头像更新成功", "avatar_url": current_user.avatar})
        
        return jsonify({"code": 400, "error": "无效的图片数据"}), 400
        
    except Exception as e:
        db.session.rollback()
        return jsonify({"code": 500, "error": str(e)}), 500


@api_bp.route('/user/stats', methods=['GET'])
@login_required
def user_stats():
    try:
        from sqlalchemy import func
        
        total_conversations = AnalysisHistory.query.filter_by(user_id=current_user.id).count()
        
        total_messages = 0
        histories = AnalysisHistory.query.filter_by(user_id=current_user.id).all()
        for h in histories:
            if h.messages:
                total_messages += len(h.messages)
        
        agent_stats = db.session.query(
            AnalysisHistory.agent_used,
            func.count(AnalysisHistory.id)
        ).filter_by(user_id=current_user.id).group_by(
            AnalysisHistory.agent_used
        ).all()
        
        agent_usage = {agent: count for agent, count in agent_stats}
        
        recent_days = db.session.query(
            func.date(AnalysisHistory.created_at).label('date'),
            func.count(AnalysisHistory.id).label('count')
        ).filter_by(user_id=current_user.id).group_by(
            func.date(AnalysisHistory.created_at)
        ).order_by(
            func.date(AnalysisHistory.created_at).desc()
        ).limit(7).all()
        
        daily_activity = [{"date": str(date), "count": count} for date, count in recent_days]
        
        return jsonify({
            "code": 200,
            "data": {
                "total_conversations": total_conversations,
                "total_messages": total_messages,
                "agent_usage": agent_usage,
                "daily_activity": daily_activity
            }
        })
    except Exception as e:
        return jsonify({"code": 500, "error": str(e)}), 500


@api_bp.route('/export/conversation/<conversation_id>', methods=['GET'])
@login_required
def export_conversation(conversation_id):
    try:
        history = AnalysisHistory.query.filter_by(
            conversation_id=conversation_id,
            user_id=current_user.id
        ).first()
        
        if not history:
            return jsonify({"code": 404, "error": "对话不存在"}), 404
        
        messages = history.messages or []
        
        text_content = f"对话标题: {history.title or '新对话'}\n"
        text_content += f"创建时间: {history.created_at.strftime('%Y-%m-%d %H:%M')}\n"
        text_content += f"使用Agent: {history.agent_used}\n"
        text_content += "=" * 50 + "\n\n"
        
        for msg in messages:
            role = "用户" if msg.get("role") == "user" else "AI助手"
            content = msg.get("content", "")
            timestamp = msg.get("timestamp", "")
            if timestamp:
                try:
                    dt = datetime.fromisoformat(timestamp.replace('Z', '+00:00'))
                    time_str = dt.strftime('%H:%M:%S')
                except:
                    time_str = ""
            else:
                time_str = ""
            
            text_content += f"[{time_str}] {role}:\n{content}\n\n"
        
        return jsonify({
            "code": 200,
            "data": {
                "title": history.title or '新对话',
                "content": text_content
            }
        })
    except Exception as e:
        return jsonify({"code": 500, "error": str(e)}), 500


@api_bp.route('/upload/resume', methods=['POST'])
@login_required
def upload_resume():
    """上传简历文件，解析并返回文本内容"""
    if 'file' not in request.files:
        return jsonify({"code": 400, "error": "没有上传文件"}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({"code": 400, "error": "未选择文件"}), 400
    
    allowed_extensions = {'.pdf', '.docx', '.doc', '.txt'}
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in allowed_extensions:
        return jsonify({"code": 400, "error": "不支持的文件格式，请上传PDF、DOCX或TXT文件"}), 400
    
    try:
        text_content = ""
        
        if ext == '.pdf':
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(file.read()))
            for page in reader.pages:
                text_content += page.extract_text() + "\n"
        
        elif ext in ('.docx', '.doc'):
            from docx import Document
            doc = Document(io.BytesIO(file.read()))
            for para in doc.paragraphs:
                text_content += para.text + "\n"
        
        elif ext == '.txt':
            text_content = file.read().decode('utf-8')
        
        if not text_content.strip():
            return jsonify({"code": 400, "error": "文件内容为空或无法解析"}), 400
        
        return jsonify({
            "code": 200,
            "data": {
                "filename": file.filename,
                "content": text_content.strip(),
                "file_type": ext
            }
        })
    
    except Exception as e:
        return jsonify({"code": 500, "error": f"文件解析失败: {str(e)}"}), 500


@api_bp.route('/download/resume', methods=['POST'])
@login_required
def download_resume():
    """将简历内容转换为文件下载（支持HTML和DOCX格式）"""
    data = request.get_json()
    content = data.get('content', '')
    filename = data.get('filename', 'resume')
    file_format = data.get('format', 'docx')
    
    if not content:
        return jsonify({"code": 400, "error": "内容不能为空"}), 400
    
    try:
        # 检测是否为HTML格式
        is_html = content.strip().startswith('<!DOCTYPE html>') or content.strip().startswith('<html')
        
        if file_format == 'html' or (file_format == 'docx' and is_html):
            # 如果请求HTML格式，或者内容是HTML但请求DOCX（先尝试HTML导出）
            if file_format == 'html':
                # 直接导出HTML文件
                buffer = io.BytesIO(content.encode('utf-8'))
                return send_file(
                    buffer,
                    as_attachment=True,
                    download_name=f'{filename}.html',
                    mimetype='text/html'
                )
            else:
                # HTML内容转DOCX：提取文本内容
                from docx import Document
                from docx.shared import Pt, Inches, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                
                doc = Document()
                
                style = doc.styles['Normal']
                style.font.name = 'Arial'
                style.font.size = Pt(11)
                
                # 简单提取HTML中的文本内容
                import re
                # 移除script和style标签
                text_content = re.sub(r'<script[^>]*>.*?</script>', '', content, flags=re.DOTALL)
                text_content = re.sub(r'<style[^>]*>.*?</style>', '', text_content, flags=re.DOTALL)
                # 移除HTML标签，保留文本
                text_content = re.sub(r'<[^>]+>', '\n', text_content)
                # 清理多余空白
                text_content = re.sub(r'\n\s*\n', '\n\n', text_content)
                text_content = text_content.strip()
                
                # 按行处理
                lines = text_content.split('\n')
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                    # 简单格式化
                    if len(line) < 50 and line.isupper():
                        doc.add_heading(line, level=1)
                    elif line.endswith('：') or line.endswith(':'):
                        p = doc.add_paragraph()
                        run = p.add_run(line)
                        run.bold = True
                    else:
                        doc.add_paragraph(line)
                
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                
                return send_file(
                    buffer,
                    as_attachment=True,
                    download_name=f'{filename}.docx',
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
        
        elif file_format == 'docx':
            # Markdown格式转DOCX
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(11)
            
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    doc.add_paragraph('')
                    continue
                
                if line.startswith('# '):
                    p = doc.add_heading(line[2:].strip(), level=1)
                elif line.startswith('## '):
                    p = doc.add_heading(line[3:].strip(), level=2)
                elif line.startswith('### '):
                    p = doc.add_heading(line[4:].strip(), level=3)
                elif line.startswith('- ') or line.startswith('* '):
                    doc.add_paragraph(line[2:].strip(), style='List Bullet')
                elif line.startswith('**') and line.endswith('**'):
                    p = doc.add_paragraph()
                    run = p.add_run(line[2:-2].strip())
                    run.bold = True
                else:
                    p = doc.add_paragraph(line)
                    if '**' in line:
                        parts = line.split('**')
                        p.clear()
                        for i, part in enumerate(parts):
                            run = p.add_run(part)
                            if i % 2 == 1:
                                run.bold = True
            
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            return send_file(
                buffer,
                as_attachment=True,
                download_name=f'{filename}.docx',
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        
        else:
            return jsonify({"code": 400, "error": "暂只支持DOCX和HTML格式"}), 400
    
    except Exception as e:
        return jsonify({"code": 500, "error": f"生成文件失败: {str(e)}"}), 500


@api_bp.route('/profile/completion', methods=['GET'])
@login_required
def get_profile_completion():
    """计算档案完善度"""
    from app.models.profile import UserProfile
    
    try:
        db.session.rollback()
    except Exception:
        pass
    
    profile = UserProfile.query.filter_by(user_id=current_user.id).first()
    
    # 字段权重
    field_weights = {
        'education': 5,
        'major': 5,
        'work_experience': 5,
        'current_job_title': 5,
        'skills': 10,
        'target_job_title': 10,
        'target_industry': 5,
        'target_salary_min': 5,
        'location_preference': 5,
        'job_search_status': 5,
        'work_preference': 5,
        'company_type_preference': 5,
        'projects': 10,
        'certifications': 5,
        'career_goals': 5,
    }
    
    total = sum(field_weights.values())
    filled = 0
    missing_fields = []
    
    if profile:
        for field, weight in field_weights.items():
            value = getattr(profile, field, None)
            is_filled = False
            
            if value is not None:
                if isinstance(value, list) and len(value) > 0:
                    is_filled = True
                elif isinstance(value, str) and value.strip():
                    is_filled = True
                elif isinstance(value, (int, float)) and value > 0:
                    is_filled = True
            
            if is_filled:
                filled += weight
            else:
                missing_fields.append(field)
    
    completion = int(filled / total * 100)
    
    return jsonify({
        "code": 200,
        "data": {
            "completion": completion,
            "filled": filled,
            "total": total,
            "missing_fields": missing_fields
        }
    })


@api_bp.route('/profile/milestones', methods=['GET'])
@login_required
def get_milestones():
    """获取决策里程碑（对话成果）"""
    try:
        db.session.rollback()
    except Exception:
        pass
    
    histories = AnalysisHistory.query.filter_by(
        user_id=current_user.id
    ).order_by(AnalysisHistory.updated_at.desc()).limit(20).all()
    
    milestones = []
    for h in histories:
        achievements = _extract_achievements(h)
        if achievements:
            milestones.append({
                "conversation_id": h.conversation_id,
                "date": h.updated_at.strftime('%m-%d') if h.updated_at else h.created_at.strftime('%m-%d'),
                "agent": h.agent_used,
                "title": achievements[0],
                "achievements": achievements
            })
    
    return jsonify({
        "code": 200,
        "data": milestones
    })


@api_bp.route('/profile/next-actions', methods=['GET'])
@login_required
def get_next_actions():
    """获取Next Action动态建议"""
    from app.models.profile import UserProfile
    
    try:
        db.session.rollback()
    except Exception:
        pass
    
    profile = UserProfile.query.filter_by(user_id=current_user.id).first()
    
    # 获取对话历史
    histories = AnalysisHistory.query.filter_by(
        user_id=current_user.id
    ).order_by(AnalysisHistory.updated_at.desc()).limit(20).all()
    
    # 分析用户状态
    actions = []
    used_agents = set()
    has_resume = False
    has_target_job = False
    has_skill_analysis = False
    has_interview_prep = False
    
    # 分析对话历史中的成就
    for h in histories:
        agent = h.agent_used or ''
        used_agents.add(agent)
        output = (h.result_data or {}).get('output', '')
        
        if agent == 'resume' or '简历' in output:
            if 'RESUME_START' in output or '个人简介' in output:
                has_resume = True
        if agent == 'skill' or '差距' in output or '学习路径' in output:
            has_skill_analysis = True
        if agent == 'interview' or '面试题' in output or '自我介绍' in output:
            has_interview_prep = True
    
    # 检查档案状态
    if profile:
        has_target_job = bool(profile.target_job_title)
        has_skills = bool(profile.skills and len(profile.skills) > 0)
        has_projects = bool(profile.projects and len(profile.projects) > 0)
        has_career_goals = bool(profile.career_goals and profile.career_goals.strip())
        profile_completion = _calculate_completion(profile)
    else:
        has_skills = False
        has_projects = False
        has_career_goals = False
        profile_completion = 0
    
    # 生成建议行动（按优先级排序）
    
    # 1. 完善档案（如果完善度低于50%）
    if profile_completion < 50:
        actions.append({
            "id": "complete_profile",
            "title": "完善职业档案",
            "desc": f"档案完善度仅{profile_completion}%，完善后AI能提供更精准的建议",
            "icon": "profile",
            "color": "#0ea5e9",
            "action": "navigate",
            "target": "#careerProfile",
            "priority": 1
        })
    
    # 2. 确定目标岗位
    if not has_target_job:
        actions.append({
            "id": "set_target_job",
            "title": "确定目标岗位",
            "desc": "设定目标岗位后，AI可以为你分析技能差距和生成简历",
            "icon": "target",
            "color": "#8b5cf6",
            "action": "chat",
            "target": "帮我分析一下适合我的目标岗位",
            "priority": 2
        })
    
    # 3. 添加技能
    if not has_skills:
        actions.append({
            "id": "add_skills",
            "title": "添加技能标签",
            "desc": "添加你的技能，让AI更好地匹配职位和分析差距",
            "icon": "skill",
            "color": "#10b981",
            "action": "navigate",
            "target": "#careerProfile",
            "priority": 3
        })
    
    # 4. 生成简历
    if has_target_job and not has_resume:
        actions.append({
            "id": "generate_resume",
            "title": "生成专业简历",
            "desc": f"你已设定目标岗位「{profile.target_job_title}」，可以生成针对性简历",
            "icon": "resume",
            "color": "#f59e0b",
            "action": "chat",
            "target": "帮我生成简历",
            "priority": 4
        })
    
    # 5. 技能差距分析
    if has_target_job and not has_skill_analysis:
        actions.append({
            "id": "skill_analysis",
            "title": "分析技能差距",
            "desc": "了解目标岗位需要的技能，制定学习计划",
            "icon": "chart",
            "color": "#ec4899",
            "action": "chat",
            "target": "分析我的技能与目标岗位的差距",
            "priority": 5
        })
    
    # 6. 面试准备
    if has_target_job and has_resume and not has_interview_prep:
        actions.append({
            "id": "interview_prep",
            "title": "准备面试",
            "desc": "获取面试题和自我介绍优化建议",
            "icon": "interview",
            "color": "#6366f1",
            "action": "chat",
            "target": "帮我准备面试",
            "priority": 6
        })
    
    # 7. 探索副业（如果有空闲时间）
    if profile and profile.available_hours_per_week and profile.available_hours_per_week >= 5:
        if 'side_job' not in used_agents:
            actions.append({
                "id": "explore_side_job",
                "title": "探索副业机会",
                "desc": f"你每周有{profile.available_hours_per_week}小时可用，探索副业增加收入",
                "icon": "money",
                "color": "#14b8a6",
                "action": "chat",
                "target": "帮我推荐适合的副业",
                "priority": 7
            })
    
    # 8. 添加项目经历
    if not has_projects and profile_completion >= 50:
        actions.append({
            "id": "add_projects",
            "title": "添加项目经历",
            "desc": "添加项目经历可以让简历更有竞争力",
            "icon": "project",
            "color": "#0ea5e9",
            "action": "navigate",
            "target": "#careerProfile",
            "priority": 8
        })
    
    # 9. 设定职业目标
    if not has_career_goals and profile_completion >= 50:
        actions.append({
            "id": "set_career_goals",
            "title": "设定职业目标",
            "desc": "明确职业目标，让AI更好地规划你的职业路径",
            "icon": "goal",
            "color": "#8b5cf6",
            "action": "navigate",
            "target": "#careerProfile",
            "priority": 9
        })
    
    # 10. 默认行动
    if not actions:
        actions.append({
            "id": "start_chat",
            "title": "开始新对话",
            "desc": "与AI顾问聊聊你的职业困惑",
            "icon": "chat",
            "color": "#0ea5e9",
            "action": "chat",
            "target": "",
            "priority": 10
        })
    
    # 按优先级排序，最多返回4个
    actions.sort(key=lambda x: x['priority'])
    actions = actions[:4]
    
    return jsonify({
        "code": 200,
        "data": actions
    })


def _calculate_completion(profile):
    """计算档案完善度"""
    field_weights = {
        'education': 5,
        'major': 5,
        'work_experience': 5,
        'current_job_title': 5,
        'skills': 10,
        'target_job_title': 10,
        'target_industry': 5,
        'target_salary_min': 5,
        'location_preference': 5,
        'job_search_status': 5,
        'work_preference': 5,
        'company_type_preference': 5,
        'projects': 10,
        'certifications': 5,
        'career_goals': 5,
    }
    
    total = sum(field_weights.values())
    filled = 0
    
    for field, weight in field_weights.items():
        value = getattr(profile, field, None)
        is_filled = False
        
        if value is not None:
            if isinstance(value, list) and len(value) > 0:
                is_filled = True
            elif isinstance(value, str) and value.strip():
                is_filled = True
            elif isinstance(value, (int, float)) and value > 0:
                is_filled = True
        
        if is_filled:
            filled += weight
    
    return int(filled / total * 100)


@api_bp.route('/profile/update', methods=['POST'])
@login_required
def update_profile():
    """更新用户档案字段"""
    from app.models.profile import UserProfile
    
    data = request.get_json()
    field = data.get('field')
    value = data.get('value')
    action = data.get('action', 'set')  # set, add, remove
    
    if not field:
        return jsonify({"code": 400, "error": "缺少field参数"}), 400
    
    try:
        db.session.rollback()
    except Exception:
        pass
    
    profile = UserProfile.query.filter_by(user_id=current_user.id).first()
    if not profile:
        profile = UserProfile(user_id=current_user.id)
        db.session.add(profile)
    
    # 字段映射
    field_map = {
        'target_job_title': 'target_job_title',
        'target_salary_min': 'target_salary_min',
        'target_salary_max': 'target_salary_max',
        'location_preference': 'location_preference',
        'work_preference': 'work_preference',
        'skills': 'skills',
    }
    
    db_field = field_map.get(field)
    if not db_field:
        return jsonify({"code": 400, "error": f"不支持的字段: {field}"}), 400
    
    try:
        if field == 'skills':
            # 技能特殊处理
            current_skills = profile.skills or []
            if action == 'add' and value:
                if value not in current_skills:
                    current_skills.append(value)
            elif action == 'remove' and value:
                current_skills = [s for s in current_skills if s != value]
            elif action == 'set':
                current_skills = value if isinstance(value, list) else [value]
            profile.skills = current_skills
            flag_modified(profile, 'skills')
        elif 'salary' in field:
            # 薪资转数字
            try:
                setattr(profile, db_field, float(value) if value else None)
            except ValueError:
                return jsonify({"code": 400, "error": "薪资必须是数字"}), 400
        else:
            setattr(profile, db_field, value if value else None)
        
        profile.updated_at = datetime.utcnow()
        db.session.commit()
        
        return jsonify({"code": 200, "message": "更新成功"})
        
    except Exception as e:
        db.session.rollback()
        return jsonify({"code": 500, "error": str(e)}), 500


def _extract_achievements(history):
    """从对话记录中提取成果标签"""
    achievements = []
    agent = history.agent_used or ''
    output = (history.result_data or {}).get('output', '')
    
    if agent == 'resume' or '简历' in output:
        if 'RESUME_START' in output or '个人简介' in output:
            achievements.append('生成了简历')
    if agent == 'skill' or '技能' in output:
        if '差距' in output or '学习路径' in output:
            achievements.append('完成了技能分析')
    if agent == 'career' or '职位' in output:
        if '目标岗位' in output or 'save_target_job' in str(history.reasoning_steps):
            achievements.append('确定了目标岗位')
        elif '薪资' in output:
            achievements.append('查询了薪资')
        elif '搜索' in output or '职位' in output:
            achievements.append('搜索了职位')
    if agent == 'interview' or '面试' in output:
        if '面试题' in output or '自我介绍' in output:
            achievements.append('获取了面试指导')
    if agent == 'side_job' or '副业' in output:
        if '推荐' in output or '匹配' in output:
            achievements.append('探索了副业机会')
    
    return achievements if achievements else [history.title or '进行了对话']
