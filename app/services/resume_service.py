"""
简历服务模块
处理简历上传下载相关的业务逻辑
"""

import io
import os
import re
from flask import send_file


class ResumeService:
    """简历服务类"""
    
    # 缓存模型列表，避免每次都调用API
    _cached_models = None
    _cache_timestamp = 0
    
    @staticmethod
    def _compress_image(image_bytes, max_size=800):
        """压缩图像，减少base64数据量"""
        try:
            from PIL import Image
            import io
            
            img = Image.open(io.BytesIO(image_bytes))
            
            # 转换为RGB（如果是RGBA）
            if img.mode in ('RGBA', 'P'):
                img = img.convert('RGB')
            
            # 调整大小
            if max(img.size) > max_size:
                img.thumbnail((max_size, max_size), Image.Resampling.LANCZOS)
            
            # 压缩为JPEG
            output = io.BytesIO()
            img.save(output, format='JPEG', quality=80, optimize=True)
            return output.getvalue()
        except Exception:
            # 压缩失败，返回原图
            return image_bytes
    
    @staticmethod
    def _extract_text_from_image_bytes(image_bytes, mimetype, filename):
        """调用LLM Vision接口提取图片中的招聘岗位信息或文字内容（优化版）"""
        from flask import current_app
        from openai import OpenAI
        import base64
        import time
        
        api_key = current_app.config.get('VISION_API_KEY') or current_app.config['OPENAI_API_KEY']
        base_url = current_app.config.get('VISION_BASE_URL') or current_app.config['OPENAI_BASE_URL']
        
        client = OpenAI(
            api_key=api_key,
            base_url=base_url
        )
        
        # 压缩图像（减少数据量）
        compressed_bytes = ResumeService._compress_image(image_bytes)
        img_base64 = base64.b64encode(compressed_bytes).decode('utf-8')
        
        if not mimetype or 'image' not in mimetype:
            mimetype = 'image/jpeg'
            
        # 精简prompt（减少token消耗）
        prompt = """提取图片中的文字内容。
如果是招聘岗位，整理出：岗位名称、职责、要求、薪资、地点。
如果是简历，提取个人信息、工作经历、技能。
只输出提取的文本，不要解释。"""

        # 直接使用配置的模型，不获取模型列表（节省一次API调用）
        configured_model = current_app.config.get('VISION_MODEL') or current_app.config['OPENAI_MODEL']
        
        # 只尝试1-2个模型，不串行尝试所有
        models_to_try = [configured_model]
        if configured_model != 'gpt-4o-mini':
            models_to_try.append('gpt-4o-mini')  # 快速兜底
                
        last_err = None
        for model_name in models_to_try:
            if not model_name:
                continue
            try:
                response = client.chat.completions.create(
                    model=model_name,
                    messages=[
                        {
                            "role": "user",
                            "content": [
                                {"type": "text", "text": prompt},
                                {
                                    "type": "image_url",
                                    "image_url": {
                                        "url": f"data:{mimetype};base64,{img_base64}"
                                    }
                                }
                            ]
                        }
                    ],
                    max_tokens=1000,  # 减少max_tokens
                    temperature=0.1
                )
                return response.choices[0].message.content.strip()
            except Exception as e:
                last_err = e
                continue
                
        raise RuntimeError(
            f"图像解析失败: {str(last_err)}"
        )

    @staticmethod
    def upload_resume(file):
        """上传文件（简历/图片/文档），解析并返回文本内容"""
        if not file or file.filename == '':
            return {"success": False, "error": "未选择文件"}
        
        allowed_extensions = {'.pdf', '.docx', '.doc', '.txt', '.png', '.jpg', '.jpeg', '.webp', '.bmp'}
        ext = os.path.splitext(file.filename)[1].lower()
        if ext not in allowed_extensions:
            return {"success": False, "error": "不支持的文件格式，请上传PDF、DOCX、TXT或图片文件"}
        
        try:
            text_content = ""
            file_bytes = file.read()
            if not file_bytes:
                return {"success": False, "error": "文件内容为空"}
            
            is_image = ext in {'.png', '.jpg', '.jpeg', '.webp', '.bmp'}
            
            if is_image:
                text_content = ResumeService._extract_text_from_image_bytes(file_bytes, file.content_type, file.filename)
            
            elif ext == '.pdf':
                from pypdf import PdfReader
                reader = PdfReader(io.BytesIO(file_bytes))

                # 安全限制：PDF 最多 50 页，防止 DoS
                if len(reader.pages) > 50:
                    return {"success": False, "error": "PDF 文件页数过多（最多支持50页）"}

                # 1. 尝试标准文本提取
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n"
                
                # 2. 如果文本为空，尝试提取表单字段与注释
                if not text_content.strip():
                    fields = reader.get_fields()
                    if fields:
                        for field_name, field_val in fields.items():
                            if isinstance(field_val, dict) and field_val.get('/V'):
                                text_content += f"{field_name}: {field_val['/V']}\n"
                            elif hasattr(field_val, 'get') and field_val.get('/V'):
                                text_content += f"{field_name}: {field_val.get('/V')}\n"
                    
                    for page in reader.pages:
                        if '/Annots' in page:
                            for annot in page['/Annots']:
                                try:
                                    annot_obj = annot.get_object()
                                    if '/Contents' in annot_obj:
                                        text_content += str(annot_obj['/Contents']) + "\n"
                                except Exception:
                                    continue
                
                # 3. 如果文本仍然为空，尝试使用大模型 Vision OCR 解析 PDF 中的图片 (扫描件 PDF)
                if not text_content.strip():
                    extracted_images_text = ""
                    image_count = 0
                    for page_idx, page in enumerate(reader.pages):
                        page_images = list(page.images)
                        for img_idx, img in enumerate(page_images[:3]):
                            try:
                                img_bytes = img.data
                                img_name = img.name or f"page_{page_idx}_img_{img_idx}"
                                img_ext = os.path.splitext(img_name)[1].lower()
                                img_mimetype = "image/png" if img_ext == ".png" else "image/jpeg"
                                
                                img_text = ResumeService._extract_text_from_image_bytes(
                                    img_bytes, 
                                    img_mimetype, 
                                    f"PDF_Page_{page_idx+1}_Image_{img_idx+1}"
                                )
                                if img_text:
                                    extracted_images_text += f"\n--- [PDF第 {page_idx+1} 页提取内容] ---\n{img_text}\n"
                                    image_count += 1
                                    if image_count >= 5: # 最多处理 5 张图片，防止大模型接口超时
                                        break
                            except Exception:
                                continue
                        if image_count >= 5:
                            break
                    text_content = extracted_images_text
            
            elif ext in ('.docx', '.doc'):
                from docx import Document
                doc = Document(io.BytesIO(file_bytes))
                for para in doc.paragraphs:
                    text_content += para.text + "\n"
            
            elif ext == '.txt':
                text_content = file_bytes.decode('utf-8')
            
            if not text_content.strip():
                return {"success": False, "error": "文件内容为空或无法解析，请确保文件包含文字或清晰可读的截图"}
            
            return {
                "success": True,
                "data": {
                    "filename": file.filename,
                    "content": text_content.strip(),
                    "file_type": ext
                }
            }
        
        except Exception as e:
            return {"success": False, "error": f"文件解析失败: {str(e)}"}
    
    @staticmethod
    def download_resume(content, filename, file_format):
        """将简历内容转换为文件下载"""
        if not content:
            return {"success": False, "error": "内容不能为空"}
        
        try:
            # 检测是否为HTML格式
            is_html = content.strip().startswith('<!DOCTYPE html>') or content.strip().startswith('<html')
            
            if file_format == 'html':
                # 直接导出HTML文件
                buffer = io.BytesIO(content.encode('utf-8'))
                return {
                    "success": True,
                    "file": send_file(
                        buffer,
                        as_attachment=True,
                        download_name=f'{filename}.html',
                        mimetype='text/html'
                    )
                }
            
            elif file_format == 'pdf':
                # 生成PDF格式 - 使用xhtml2pdf（纯Python，无需系统依赖）
                from xhtml2pdf import pisa
                
                if is_html:
                    html_content = content
                else:
                    # Markdown转HTML，保留基本样式
                    html_content = content
                    html_content = re.sub(
                        r'^# (.+)$', 
                        r'<h1 style="font-size:20px;color:#333;border-bottom:2px solid #4285f4;padding-bottom:8px;">\1</h1>', 
                        html_content, flags=re.MULTILINE
                    )
                    html_content = re.sub(
                        r'^## (.+)$', 
                        r'<h2 style="font-size:16px;color:#4285f4;margin-top:20px;">\1</h2>', 
                        html_content, flags=re.MULTILINE
                    )
                    html_content = re.sub(
                        r'^### (.+)$', 
                        r'<h3 style="font-size:14px;color:#555;">\1</h3>', 
                        html_content, flags=re.MULTILINE
                    )
                    html_content = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', html_content)
                    html_content = re.sub(
                        r'^- (.+)$', 
                        r'<li style="margin:4px 0;">\1</li>', 
                        html_content, flags=re.MULTILINE
                    )
                    html_content = re.sub(
                        r'(<li.*?</li>)', 
                        r'<ul style="padding-left:20px;">\1</ul>', 
                        html_content, flags=re.DOTALL
                    )
                    html_content = html_content.replace('\n', '<br>')
                    html_content = f'''<!DOCTYPE html>
<html><head><meta charset="utf-8"></head>
<body style="font-family:'Microsoft YaHei','SimSun',Arial,sans-serif;font-size:12pt;line-height:1.8;padding:20px;color:#333;">
{html_content}
</body></html>'''
                
                # 确保HTML包含中文字体声明
                if '<style' not in html_content:
                    # 在head中添加字体样式
                    html_content = html_content.replace(
                        '<head>',
                        '<head><style>@page { size: A4; margin: 2cm; } body { font-family: "Microsoft YaHei", "SimSun", "STSong", Arial, sans-serif; }</style>'
                    )
                
                buffer = io.BytesIO()
                pisa_status = pisa.CreatePDF(html_content, dest=buffer, encoding='utf-8')
                
                if pisa_status.err:
                    return {"success": False, "error": "PDF生成失败"}
                
                buffer.seek(0)
                
                return {
                    "success": True,
                    "file": send_file(
                        buffer,
                        as_attachment=True,
                        download_name=f'{filename}.pdf',
                        mimetype='application/pdf'
                    )
                }
            
            elif file_format == 'docx':
                # Markdown格式转DOCX
                from docx import Document
                from docx.shared import Pt
                
                doc = Document()
                
                style = doc.styles['Normal']
                style.font.name = 'Arial'
                style.font.size = Pt(11)
                
                lines = content.split('\n')
                for line in lines:
                    line = line.strip()
                    if not line:
                        doc.add_paragraph('')
                        continue
                    
                    if line.startswith('# '):
                        doc.add_heading(line[2:].strip(), level=1)
                    elif line.startswith('## '):
                        doc.add_heading(line[3:].strip(), level=2)
                    elif line.startswith('### '):
                        doc.add_heading(line[4:].strip(), level=3)
                    elif line.startswith('- ') or line.startswith('* '):
                        doc.add_paragraph(line[2:].strip(), style='List Bullet')
                    elif line.startswith('**') and line.endswith('**'):
                        p = doc.add_paragraph()
                        run = p.add_run(line[2:-2].strip())
                        run.bold = True
                    else:
                        p = doc.add_paragraph(line)
                        if '**' in line:
                            parts = line.split('**')
                            p.clear()
                            for i, part in enumerate(parts):
                                run = p.add_run(part)
                                if i % 2 == 1:
                                    run.bold = True
                
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                
                return {
                    "success": True,
                    "file": send_file(
                        buffer,
                        as_attachment=True,
                        download_name=f'{filename}.docx',
                        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                    )
                }
            
            else:
                return {"success": False, "error": "暂只支持HTML、PDF和DOCX格式"}
        
        except Exception as e:
            return {"success": False, "error": f"生成文件失败: {str(e)}"}
